import os
import docx
import pandas as pd
import fitz  # PyMuPDF

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"❌ خطأ أثناء فتح ملف الوورد: {e}")
        return []

    full_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    full_text.append(text)

    return full_text

def extract_text_from_excel(file_path):
    try:
        xls = pd.ExcelFile(file_path)
    except Exception as e:
        print(f"❌ خطأ أثناء فتح ملف الإكسل: {e}")
        return []

    all_text = []

    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str)

        all_text.append(f"--- Sheet: {sheet_name} ---")

        for row in df.itertuples(index=False):
            for cell in row:
                if pd.notna(cell):
                    text = str(cell).strip()
                    if text:
                        all_text.append(text)

    return all_text

def extract_text_from_pdf(file_path):
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        print(f"❌ خطأ أثناء فتح ملف PDF: {e}")
        return []

    full_text = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text().strip()
        if text:
            full_text.append(f"--- Page {page_num + 1} ---")
            full_text.append(text)

    return full_text

def main():
    file_path = input("📂 Enter the path of the file (Word .docx, Excel .xlsx/.xls, or PDF .pdf): ").strip()

    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return

    ext = file_path.lower()

    if ext.endswith('.docx'):
        extracted_text = extract_text_from_docx(file_path)
    elif ext.endswith(('.xlsx', '.xls')):
        extracted_text = extract_text_from_excel(file_path)
    elif ext.endswith('.pdf'):
        extracted_text = extract_text_from_pdf(file_path)
    else:
        print("❌ Unsupported file format. Please provide a .docx, .xlsx/.xls, or .pdf file.")
        return

    if not extracted_text:
        print("⚠ No text found in the file.")
        return

    print("\n📄 Extracted text:")
    for line in extracted_text:
        print(line)

    with open("output.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(extracted_text))

    print("\n✅ Text saved to output.txt")

if __name__ == "__main__":
    main()